# -*- coding: utf-8 -*-
"""
SMUROP Document Factory (v4.0)
Based on v2.3 (proven working).

Added in v4:
- --doc-types: dosie, dnevnik, audiogram, roditelski, izvestaj (comma-sep)
- --output-mode: separate | merged | both
- --therapist-name: editable therapist name (default: Благој Насев)
- --institution: editable institution header
- Родителски средби document type (individual parent meetings per student)
- Извештај (сумирано) document type
- All text in Cyrillic (no Latin transliterations)
- Fixed Unicode console error on Windows
- Fixed font_name/font_size global scope issue
- Proper Heading 1/2/3 for PDF bookmarks

Expected inputs:
- student_db_backup_*.json  (meta + data.student_records)
- diary_export_*.json       (students + plans + trijazenTestovi + studentProgress + attendance)
- AUDIOGRAMI_BAZA_*.json    (array of {subjectName,date,rightAir,..., imageLinks?})
or
- unified JSON (e_dnevnik_unified_state*.json / similar) containing:
  students + student_records + audiograms + trijazenTestovi
"""
import argparse
import base64
import binascii
import datetime
import io
import json
import os
import re
import sys
import tempfile
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm, RGBColor

try:
    import requests
except Exception:
    requests = None


# ----------------------------
# Global config (set by main, used by generators)
# ----------------------------
class CFG:
    font_name = "Times New Roman"
    font_size = 12
    school_year = "2025-2026"
    therapist_name = "Благој Насев"
    therapist_title = "дипл. деф. сурдолог-аудиорехабилитатор"
    institution = 'Завод за рехабилитација на деца со оштетен слух, говор, глас и други проблеми во развојот "Кочо Рацин" – Битола'
    institution_short = 'ОУРЦ "Кочо Рацин" Битола'
    cabinet = "Кабинет за слух, говор, глас, алтернативна и аугментативна комуникација"


# ----------------------------
# IO helpers
# ----------------------------
def load_json(p: Path) -> Any:
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)


def save_json(p: Path, obj: Any) -> None:
    with open(p, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)


def safe_date_fmt(iso: str) -> str:
    if not iso:
        return ""
    try:
        return datetime.date.fromisoformat(iso).strftime("%d.%m.%Y")
    except Exception:
        return iso


def is_image_path(path_or_url: str) -> bool:
    s = (path_or_url or "").lower().strip()
    return any(s.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp"))


def decode_data_image_url(data_url: str) -> Optional[Tuple[str, bytes]]:
    """
    Decode data URLs like: data:image/png;base64,AAAA...
    Returns (mime_type, raw_bytes) or None.
    """
    s = (data_url or "").strip()
    if not s.lower().startswith("data:image/"):
        return None
    if "," not in s:
        return None
    header, payload = s.split(",", 1)
    if ";base64" not in header.lower():
        return None
    try:
        mime = header[5:].split(";", 1)[0].strip().lower()  # strip "data:"
    except Exception:
        return None
    if not mime.startswith("image/"):
        return None
    payload = re.sub(r"\s+", "", payload or "")
    if not payload:
        return None
    try:
        raw = base64.b64decode(payload, validate=False)
    except (ValueError, binascii.Error):
        return None
    if not raw:
        return None
    return mime, raw


def normalize_whitespace(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def split_full_name(full_name: str) -> Tuple[str, str]:
    name = normalize_whitespace(full_name)
    if not name:
        return "", ""
    parts = name.split(" ", 1)
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], parts[1]


def normalize_student_record(raw: Dict[str, Any], fallback_id: int = 0) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return {}

    rec = dict(raw)
    first = normalize_whitespace(str(rec.get("firstName", "")))
    last = normalize_whitespace(str(rec.get("lastName", "")))
    if not first and not last:
        fn, ln = split_full_name(str(rec.get("name", "")))
        first = first or fn
        last = last or ln

    rid = rec.get("id")
    if rid in (None, ""):
        rid = rec.get("studentId")
    if rid in (None, ""):
        rid = fallback_id

    links = rec.get("attachmentLinks")
    if not isinstance(links, list):
        links = []

    rec["id"] = rid
    rec["firstName"] = first
    rec["lastName"] = last
    rec["grade"] = rec.get("grade", "") or ""
    rec["birthDate"] = rec.get("birthDate", "") or ""
    rec["contact"] = rec.get("contact", "") or ""
    rec["fatherName"] = rec.get("fatherName", "") or ""
    rec["motherName"] = rec.get("motherName", "") or ""
    rec["address"] = rec.get("address", "") or ""
    rec["residence"] = rec.get("residence", "") or ""
    rec["findings"] = rec.get("findings", "") or ""
    rec["opinion"] = rec.get("opinion", "") or ""
    rec["attachmentLinks"] = links
    return rec


def get_student_records_ref(payload: Any) -> List[Dict[str, Any]]:
    """
    Return a mutable reference to the underlying student-record list when possible.
    """
    if isinstance(payload, dict):
        data = payload.get("data")
        if isinstance(data, dict) and isinstance(data.get("student_records"), list):
            return data["student_records"]
        if isinstance(payload.get("student_records"), list):
            return payload["student_records"]
        if isinstance(payload.get("studentRecords"), list):
            return payload["studentRecords"]
        if isinstance(payload.get("students"), list):
            return payload["students"]
    if isinstance(payload, list):
        return payload
    return []


def extract_student_records(payload: Any) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []

    if isinstance(payload, dict):
        candidates = []
        if isinstance(payload.get("data"), dict) and isinstance(payload["data"].get("student_records"), list):
            candidates = payload["data"]["student_records"]
        elif isinstance(payload.get("student_records"), list):
            candidates = payload["student_records"]
        elif isinstance(payload.get("studentRecords"), list):
            candidates = payload["studentRecords"]
        elif isinstance(payload.get("students"), list):
            candidates = payload["students"]

        for i, item in enumerate(candidates, start=1):
            rec = normalize_student_record(item if isinstance(item, dict) else {}, fallback_id=i)
            if rec:
                records.append(rec)
        return records

    if isinstance(payload, list):
        for i, item in enumerate(payload, start=1):
            rec = normalize_student_record(item if isinstance(item, dict) else {}, fallback_id=i)
            if rec:
                records.append(rec)
        return records

    return []


def extract_diary_payload(payload: Any) -> Dict[str, Any]:
    if not isinstance(payload, dict):
        return {}

    out = dict(payload)
    if not isinstance(out.get("students"), list):
        out["students"] = []
    if not isinstance(out.get("plans"), list):
        out["plans"] = []
    if not isinstance(out.get("trijazenTestovi"), list):
        out["trijazenTestovi"] = []
    if not isinstance(out.get("attendance"), dict):
        out["attendance"] = {}
    if not isinstance(out.get("studentProgress"), dict):
        out["studentProgress"] = {}
    if not isinstance(out.get("schedule"), dict):
        out["schedule"] = {}
    if not isinstance(out.get("scheduleHistory"), dict):
        out["scheduleHistory"] = {}
    if not isinstance(out.get("links"), list):
        out["links"] = []
    return out


def extract_audi_records(payload: Any) -> List[Dict[str, Any]]:
    if payload is None:
        return []

    # Already list-like payload
    if isinstance(payload, list):
        return normalize_audi_db(payload)

    if not isinstance(payload, dict):
        return []

    # Common keys
    for k in ("audiograms", "audiogramRecords"):
        if isinstance(payload.get(k), list):
            return normalize_audi_db(payload.get(k))
    if isinstance(payload.get("data"), dict) and isinstance(payload["data"].get("audiograms"), list):
        return normalize_audi_db(payload["data"].get("audiograms"))

    out: List[Dict[str, Any]] = []

    # Standalone audiogram backup format with numeric keys
    for key, value in payload.items():
        if re.fullmatch(r"\d+", str(key)) and isinstance(value, dict) and value.get("subjectName"):
            out.append({
                "subjectName": value.get("subjectName", ""),
                "date": value.get("date"),
                "rightAir": value.get("rightAir") or {},
                "rightBone": value.get("rightBone") or {},
                "leftAir": value.get("leftAir") or {},
                "leftBone": value.get("leftBone") or {},
                "recordType": value.get("recordType", "history"),
                "imageLinks": value.get("imageLinks") if isinstance(value.get("imageLinks"), list) else [],
            })

    if payload.get("subjectName") and any(payload.get(x) is not None for x in ("rightAir", "leftAir", "rightBone", "leftBone")):
        out.append({
            "subjectName": payload.get("subjectName", ""),
            "date": payload.get("date"),
            "rightAir": payload.get("rightAir") or {},
            "rightBone": payload.get("rightBone") or {},
            "leftAir": payload.get("leftAir") or {},
            "leftBone": payload.get("leftBone") or {},
            "recordType": "current_snapshot",
            "imageLinks": payload.get("imageLinks") if isinstance(payload.get("imageLinks"), list) else [],
        })

    # Unified nested students[].audiograms
    students = payload.get("students")
    if isinstance(students, list):
        for st in students:
            if not isinstance(st, dict):
                continue
            full_name = normalize_whitespace(str(st.get("name", "")).strip())
            if not full_name:
                full_name = normalize_whitespace(f"{st.get('firstName','')} {st.get('lastName','')}")
            arr = st.get("audiograms")
            if not isinstance(arr, list):
                continue
            for a in arr:
                if not isinstance(a, dict):
                    continue
                out.append({
                    "subjectName": normalize_whitespace(a.get("subjectName") or full_name),
                    "date": a.get("date"),
                    "rightAir": a.get("rightAir") or {},
                    "rightBone": a.get("rightBone") or {},
                    "leftAir": a.get("leftAir") or {},
                    "leftBone": a.get("leftBone") or {},
                    "recordType": a.get("recordType", "history"),
                    "imageLinks": a.get("imageLinks") if isinstance(a.get("imageLinks"), list) else [],
                })

    return normalize_audi_db(out)


# ----------------------------
# Audiogram DB normalization
# ----------------------------

def normalize_audi_db(obj: Any) -> List[Dict[str, Any]]:
    """Return a clean list of audiogram dict records.

    Supports:
      - list of dicts (normal)
      - list with some stray non-dict entries (they are ignored)
      - dict wrappers containing a list under common keys (data, records, audiograms)
      - list of JSON strings (each string parsed if possible)
    """
    if obj is None:
        return []

    # Wrapper object
    if isinstance(obj, dict):
        for key in ("data", "records", "audiograms", "items"):
            v = obj.get(key)
            if isinstance(v, list):
                obj = v
                break
        else:
            return []

    out: List[Dict[str, Any]] = []
    if isinstance(obj, list):
        for x in obj:
            if isinstance(x, dict):
                out.append(x)
                continue
            if isinstance(x, str):
                s = x.strip()
                # Sometimes exports accidentally store records as JSON strings
                if s.startswith("{") and s.endswith("}"):
                    try:
                        y = json.loads(s)
                        if isinstance(y, dict):
                            out.append(y)
                    except Exception:
                        pass
                continue
            # ignore any other types
        return out

    return []


# ----------------------------
# DOCX helpers
# ----------------------------
def mk_header(doc: Document) -> None:
    p = doc.add_paragraph(CFG.institution)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = CFG.font_name
    doc.add_paragraph("")

# ----------------------------
# Styling helpers
# ----------------------------

def apply_default_styles(doc: Document, font_name: str = "Times New Roman", font_size: int = 12) -> None:
    """Force Times New Roman + font size for Normal + Heading styles (Cyrillic-safe).

    Note: Headings remain headings (for PDF bookmarks) but we override color/size to match your preference.
    """
    def set_style(sty):
        if not sty:
            return
        try:
            f = sty.font
            f.name = font_name
            f.size = Pt(int(font_size))
            f.color.rgb = RGBColor(0, 0, 0)
            # Ensure Cyrillic uses same font
            rPr = sty.element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
        except Exception:
            pass

    styles = doc.styles
    # Normal
    try:
        set_style(styles["Normal"])
    except Exception:
        pass
    # Headings used for bookmarks
    for i in range(1, 4):
        try:
            set_style(styles[f"Heading {i}"])
        except Exception:
            pass




def add_section_title(doc: Document, title: str, level: int = 2) -> None:
    """Add section title as a Word Heading so PDF export can create bookmarks."""
    level = max(1, min(int(level), 9))
    style = f"Heading {level}"
    doc.add_paragraph(title, style=style)


def add_kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = "" if v is None else str(v)
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(12)
    doc.add_paragraph("")


def add_hyperlink(paragraph, url: str, text: Optional[str] = None):
    """
    Create a clickable hyperlink in a python-docx paragraph.
    """
    text = text or url

    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Basic link styling (blue + underline)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ----------------------------
# Patch logic
# ----------------------------
STUDENT_PATCH_FIELDS = [
    "birthDate",
    "fatherName",
    "motherName",
    "contact",
    "address",
    "residence",
    "findings",
    "opinion",
    "attachmentLinks",
]


def apply_student_patch(student_db: Dict[str, Any], patch: Dict[str, Any]) -> int:
    """
    patch format:
      { "<studentId>": { "birthDate": "...", ... } }
    Only non-empty values are applied.
    """
    source_records = get_student_records_ref(student_db)
    by_id: Dict[str, Dict[str, Any]] = {}
    for i, raw in enumerate(source_records, start=1):
        if not isinstance(raw, dict):
            continue
        # normalize key fields in-place so patching persists in original structure
        if raw.get("id") in (None, ""):
            raw["id"] = raw.get("studentId") or i
        if not (raw.get("firstName") or "").strip() and not (raw.get("lastName") or "").strip():
            fn, ln = split_full_name(str(raw.get("name", "")))
            if fn and not (raw.get("firstName") or "").strip():
                raw["firstName"] = fn
            if ln and not (raw.get("lastName") or "").strip():
                raw["lastName"] = ln
        if not isinstance(raw.get("attachmentLinks"), list):
            raw["attachmentLinks"] = []
        by_id[str(raw.get("id"))] = raw
    applied = 0

    for sid, changes in (patch or {}).items():
        if sid not in by_id or not isinstance(changes, dict):
            continue
        st = by_id[sid]
        changed_any = False
        for k, v in changes.items():
            if k not in STUDENT_PATCH_FIELDS:
                continue
            if v is None:
                continue
            if isinstance(v, str) and not v.strip():
                continue

            if k == "attachmentLinks":
                # allow patch to supply list; merge with existing
                if isinstance(v, list):
                    existing = st.get("attachmentLinks") or []
                    if not isinstance(existing, list):
                        existing = []
                    merged = []
                    seen = set()
                    for it in existing + v:
                        if not isinstance(it, str):
                            continue
                        t = it.strip()
                        if not t or t in seen:
                            continue
                        seen.add(t)
                        merged.append(t)
                    st["attachmentLinks"] = merged
                    changed_any = True
                continue

            st[k] = v
            changed_any = True

        if changed_any:
            applied += 1

    return applied


def parse_audi_patch(obj: Any) -> Dict[str, List[str]]:
    """
    audiogram_images_patch.json formats supported:
      1) { "Name|YYYY-MM-DD": { "imageLinks": [..] } }
      2) { "Name|YYYY-MM-DD": [..] }
    returns: {key: [links]}
    """
    out = {}
    if not isinstance(obj, dict):
        return out
    for k, v in obj.items():
        if isinstance(v, dict):
            links = v.get("imageLinks", [])
        else:
            links = v
        if not isinstance(links, list):
            continue
        clean = [x.strip() for x in links if isinstance(x, str) and x.strip()]
        if clean:
            out[str(k)] = clean
    return out


def audi_key(subject_name: str, date: str) -> str:
    return f"{normalize_whitespace(subject_name)}|{str(date or '').strip()}"


def apply_audi_image_patch(audi_db: Any, patch: Dict[str, List[str]]) -> int:
    """
    Applies imageLinks to audiogram records based on "subject|date" key.
    Overwrites imageLinks (deterministic) with patch values.
    """
    applied = 0
    for rec in normalize_audi_db(audi_db):
        k = audi_key(rec.get("subjectName", ""), rec.get("date", ""))
        if k in patch:
            rec["imageLinks"] = patch[k]
            applied += 1
    return applied


# ----------------------------
# Indexes and generation
# ----------------------------
def build_indexes(student_db: Dict[str, Any], diary: Dict[str, Any], audi_db: List[Dict[str, Any]], overrides: Dict[str, Any]):
    students = extract_student_records(student_db)
    student_names = {f'{s.get("firstName","")} {s.get("lastName","")}'.strip(): str(s.get("id")) for s in students}

    triage_by_sid = defaultdict(list)
    for t in diary.get("trijazenTestovi", []) or []:
        sid = str(t.get("studentId"))
        triage_by_sid[sid].append(t)

    audi_by_name = defaultdict(list)
    audi_list = normalize_audi_db(audi_db)
    for a in audi_list:
        name = normalize_whitespace(a.get("subjectName", ""))
        if not name:
            continue
        audi_by_name[name].append(a)

    aliases = overrides.get("aliases", {}) if isinstance(overrides, dict) else {}
    categories = overrides.get("categories", {}) if isinstance(overrides, dict) else {}
    return students, student_names, triage_by_sid, audi_by_name, aliases, categories


def try_fetch_image(url: str, timeout: int = 12) -> Optional[Path]:
    if not requests:
        return None
    try:
        r = requests.get(url, timeout=timeout)
        r.raise_for_status()
        ctype = (r.headers.get("content-type") or "").lower()
        # only save if seems image OR url has image extension
        if "image/" not in ctype and not is_image_path(url):
            return None
        suffix = ".png"
        if "jpeg" in ctype or url.lower().endswith(".jpg") or url.lower().endswith(".jpeg"):
            suffix = ".jpg"
        fd, p = tempfile.mkstemp(suffix=suffix, prefix="smurop_img_")
        os.close(fd)
        path = Path(p)
        path.write_bytes(r.content)
        return path
    except Exception:
        return None


def resolve_local_image(link: str, images_folder: Optional[Path]) -> Optional[Path]:
    """
    Supports:
    - absolute local path
    - relative filename within images_folder (search shallow + 1 level deep)
    """
    s = (link or "").strip()
    if not s:
        return None
    p = Path(s)
    if p.exists() and p.is_file():
        return p

    if images_folder:
        # try as relative
        cand = images_folder / s
        if cand.exists() and cand.is_file():
            return cand
        # try shallow search
        for found in images_folder.glob(s):
            if found.is_file():
                return found
        # one level deep
        for found in images_folder.glob(f"*/{s}"):
            if found.is_file():
                return found
    return None


def add_images_or_links(doc: Document, links: List[str], images_folder: Optional[Path], download_images: bool):
    """
    Insert pictures when possible; otherwise insert hyperlinks.
    """
    if not links:
        return
    for i, link in enumerate(links, start=1):
        link = (link or "").strip()
        if not link:
            continue

        # inline base64 data URL (always embed, independent of --download-images)
        decoded = decode_data_image_url(link)
        if decoded:
            mime, raw = decoded
            doc.add_paragraph(f"Слика {i}: inline base64 ({mime})")
            try:
                doc.add_picture(io.BytesIO(raw), width=Inches(6.3))
            except Exception:
                # keep traceability when decoding succeeds but embedding fails
                doc.add_paragraph("Base64 сликата не можеше да се вметне во DOCX.")
            doc.add_paragraph("")
            continue
        if link.lower().startswith("data:image/"):
            doc.add_paragraph(f"Слика {i}: Невалиден base64 запис.")
            doc.add_paragraph("")
            continue

        img_path = None
        # local first
        img_path = resolve_local_image(link, images_folder)

        # download if http(s)
        tmp = None
        if not img_path and download_images and link.lower().startswith(("http://", "https://")):
            tmp = try_fetch_image(link)
            img_path = tmp

        if img_path and img_path.exists() and is_image_path(str(img_path)):
            doc.add_paragraph(f"Слика {i}: {Path(link).name if not link.lower().startswith('http') else link}")
            try:
                doc.add_picture(str(img_path), width=Inches(6.3))
            except Exception:
                # fallback to hyperlink
                p = doc.add_paragraph("Линк: ")
                add_hyperlink(p, link)
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph("Линк: ")
            add_hyperlink(p, link)
            doc.add_paragraph("")

        if tmp:
            try:
                tmp.unlink(missing_ok=True)  # py>=3.8
            except Exception:
                pass


def generate_dossier(
    out_dir: Path,
    students: List[Dict[str, Any]],
    student_names: Dict[str, str],
    diary: Dict[str, Any],
    triage_by_sid,
    audi_by_name,
    aliases: Dict[str, str],
    categories: Dict[str, str],
    full_name: str,
    school_year: str = "2025-2026",
    images_folder: Optional[Path] = None,
    download_images: bool = False,
) -> Path:
    sid = student_names.get(full_name)
    if not sid:
        raise ValueError(f"Unknown student name: {full_name}")

    s = next(x for x in students if str(x.get("id")) == sid)

    # plan
    diary_student = next((x for x in (diary.get("students") or []) if str(x.get("id")) == sid), None)
    plan = None
    if diary_student:
        pid = diary_student.get("planId")
        plan = next((p for p in (diary.get("plans") or []) if p.get("id") == pid), None)

    # triage
    tri_list = sorted(triage_by_sid.get(sid, []), key=lambda x: x.get("date", ""))

    # audiograms
    name_norm = normalize_whitespace(full_name)
    name_for_audi = name_norm if name_norm in audi_by_name else normalize_whitespace(aliases.get(full_name, full_name))
    audi_list = sorted(audi_by_name.get(name_for_audi, []), key=lambda x: x.get("date", ""))

    # logs
    logs = []
    sp = (diary.get("studentProgress") or {}).get(sid, {}) or {}
    for pid, sess_list in sp.items():
        p = next((pp for pp in (diary.get("plans") or []) if str(pp.get("id")) == str(pid)), None)
        acts = p.get("activities", []) if p else []
        for sess in (sess_list or []):
            idx = sess.get("index")
            act = acts[idx] if isinstance(idx, int) and 0 <= idx < len(acts) else ""
            logs.append({"date": sess.get("date", ""), "time": sess.get("time", ""), "activity": act})
    logs.sort(key=lambda x: (x["date"], x["time"]))

    # attendance
    present = absent = 0
    for _, by_student in (diary.get("attendance") or {}).items():
        st = by_student.get(sid) if isinstance(by_student, dict) else None
        if isinstance(st, dict):
            for _, status in st.items():
                if status == "present":
                    present += 1
                elif status == "absent":
                    absent += 1
    total = present + absent
    pct = (present / total * 100) if total else None

    cat = (categories or {}).get(sid, None)
    has_triage = bool(tri_list)
    has_audi = bool(audi_list)

    if cat in ("literacy", "autism"):
        include_triage = False
        include_audi = False
    elif cat == "cochlear_implant":
        include_triage = has_triage
        include_audi = False
    elif cat in ("hearing", "speech"):
        include_triage = has_triage
        include_audi = (has_audi if cat == "hearing" else False)
    else:
        include_triage = has_triage
        include_audi = has_audi

    doc = Document()
    apply_default_styles(doc, font_name=CFG.font_name, font_size=CFG.font_size)
    mk_header(doc)

    add_section_title(doc, f"СУРДОЛОШКО/ДЕФЕКТОЛОШКО ДОСИЕ – {school_year}", level=1)
    doc.add_paragraph("")

    add_section_title(doc, "1. Општи податоци")
    add_kv_table(
        doc,
        [
            ("Ученик", full_name),
            ("Одделение", s.get("grade", "")),
            ("Датум на раѓање", safe_date_fmt(s.get("birthDate"))),
            ("Контакт", s.get("contact", "")),
            ("Татко", s.get("fatherName", "")),
            ("Мајка", s.get("motherName", "")),
            ("Адреса", s.get("address", "")),
            ("Живеалиште", s.get("residence", "")),
            ("Категорија (по потреба)", cat or "—"),
        ],
    )

    add_section_title(doc, "2. Наод и мислење")
    doc.add_paragraph(s.get("findings", "") or "—")
    doc.add_paragraph("")
    doc.add_paragraph(s.get("opinion", "") or "—")
    doc.add_paragraph("")

    add_section_title(doc, "3. План за работа")
    if plan:
        doc.add_paragraph(f"План: {plan.get('name', '')}")
        acts = plan.get("activities", []) or []
        for a in acts[:12]:
            doc.add_paragraph(str(a))
        if len(acts) > 12:
            doc.add_paragraph(f"... (вкупно {len(acts)} активности)")
    else:
        doc.add_paragraph("—")
    doc.add_paragraph("")

    add_section_title(doc, "4. Тријажен тест(ови)")
    if include_triage and tri_list:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Дата"
        table.rows[0].cells[1].text = "Област"
        table.rows[0].cells[2].text = "Ставка"
        table.rows[0].cells[3].text = "Резултат"
        for t in tri_list:
            r = table.add_row().cells
            r[0].text = safe_date_fmt(t.get("date", ""))
            r[1].text = str(t.get("area", ""))
            r[2].text = str(t.get("item", ""))
            r[3].text = str(t.get("result", ""))
    else:
        doc.add_paragraph("—")
    doc.add_paragraph("")

    if include_audi:
        add_section_title(doc, "5. Аудиограм(и)")
        for a in audi_list:
            add_section_title(doc, f"Датум: {safe_date_fmt(a.get('date'))}", level=3)
            freqs = sorted({*a.get("rightAir", {}).keys(), *a.get("leftAir", {}).keys()}, key=lambda x: int(x))
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Hz"
            table.rows[0].cells[1].text = "Десно (Air)"
            table.rows[0].cells[2].text = "Лево (Air)"
            for f in freqs:
                r = table.add_row().cells
                r[0].text = str(f)
                r[1].text = str(a.get("rightAir", {}).get(f, ""))
                r[2].text = str(a.get("leftAir", {}).get(f, ""))
            doc.add_paragraph("")

            links = a.get("imageLinks") if isinstance(a, dict) else None
            if isinstance(links, list) and links:
                doc.add_paragraph("Аудиограм – слики/линкови:")
                add_images_or_links(doc, links, images_folder, download_images)
            else:
                # nothing to add
                pass

    add_section_title(doc, "6. Дневник / реализирани часови")
    if logs:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Дата"
        table.rows[0].cells[1].text = "Време"
        table.rows[0].cells[2].text = "Што е работено"
        for lg in logs:
            r = table.add_row().cells
            r[0].text = safe_date_fmt(lg["date"])
            r[1].text = str(lg["time"] or "")
            r[2].text = str(lg["activity"] or "")
    else:
        doc.add_paragraph("—")
    doc.add_paragraph("")
    doc.add_paragraph(f"Присуство: {present} | Отсуство: {absent}" + (f" | {pct:.1f}%" if pct is not None else ""))

    add_section_title(doc, "7. Прилози (линкови/слики)")
    # First: student-level attachmentLinks (existing field in student_db_backup)
    attach_links = s.get("attachmentLinks") if isinstance(s, dict) else None
    if isinstance(attach_links, list) and attach_links:
        cleaned_links = [x.strip() for x in attach_links if isinstance(x, str) and x.strip()]
        if cleaned_links:
            doc.add_paragraph("Прикачени прилози (од досие):")
            add_images_or_links(doc, cleaned_links, images_folder, download_images)
        else:
            doc.add_paragraph("—")
            doc.add_paragraph("")
    else:
        doc.add_paragraph("—")
        doc.add_paragraph("")

    add_section_title(doc, "Потпис")
    doc.add_paragraph(f"{CFG.therapist_title} {CFG.therapist_name}: ____________________    Датум: ______________")

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"DOSIE_{full_name.replace(' ','_')}_{school_year}.docx"
    doc.save(out_path)
    return out_path


# ----------------------------
# Родителски средби generator
# ----------------------------
def generate_roditelski(
    out_dir: Path,
    students: List[Dict[str, Any]],
    student_names: Dict[str, str],
    diary: Dict[str, Any],
    school_year: str = "2025-2026",
) -> Path:
    """Generate parent meeting logs - individual meetings per student."""
    doc = Document()
    apply_default_styles(doc, font_name=CFG.font_name, font_size=CFG.font_size)

    # Title page
    mk_header(doc)
    add_section_title(doc, f"СРЕДБИ СО РОДИТЕЛИ", level=1)
    p = doc.add_paragraph(f"учебната {school_year}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"{CFG.therapist_title} {CFG.therapist_name}")
    doc.add_paragraph("")

    targets = sorted(student_names.keys(), key=lambda x: normalize_whitespace(x))

    for full_name in targets:
        sid = student_names.get(full_name)
        s = next((x for x in students if str(x.get("id")) == sid), None)
        if not s:
            continue

        add_section_title(doc, f"Ученик/чка: {full_name}", level=2)
        doc.add_paragraph("")

        # Parent names from student record
        parents = []
        father = (s.get("fatherName") or "").strip()
        mother = (s.get("motherName") or "").strip()
        if father:
            parents.append(father)
        if mother:
            parents.append(mother)

        # Generate 5 meeting boxes (Sept, Nov, Jan, Apr, Jun)
        meeting_dates = [
            ("Септември", f"{school_year[:4]}-09-15"),
            ("Ноември", f"{school_year[:4]}-11-20"),
            ("Јануари", f"{school_year[-4:]}-01-15"),
            ("Април", f"{school_year[-4:]}-04-15"),
            ("Јуни", f"{school_year[-4:]}-06-10"),
        ]

        for i, (month_label, default_date) in enumerate(meeting_dates, 1):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            # Build meeting box content
            cell_p = cell.paragraphs[0]
            run = cell_p.add_run(f"Дата: {default_date}       Родител: ")
            run.font.name = CFG.font_name
            run.font.size = Pt(CFG.font_size)
            if parents:
                parent_name = parents[i % len(parents)]
                run2 = cell_p.add_run(f"{parent_name} ({'татко' if i % 2 == 1 and father else 'мајка'} на {full_name})")
            else:
                run2 = cell_p.add_run("____________________")
            run2.font.name = CFG.font_name
            run2.font.size = Pt(CFG.font_size)

            p_tema = cell.add_paragraph()
            run_t = p_tema.add_run("Тема:")
            run_t.bold = True
            run_t.font.name = CFG.font_name
            run_t.font.size = Pt(CFG.font_size)

            # Placeholder text for meeting content
            cell.add_paragraph("")
            cell.add_paragraph("")

            p_sign = cell.add_paragraph()
            p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_s = p_sign.add_run(f"{CFG.cabinet}")
            run_s.font.name = CFG.font_name
            run_s.font.size = Pt(9)
            p_sign2 = cell.add_paragraph()
            p_sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_s2 = p_sign2.add_run(f"{CFG.therapist_title} {CFG.therapist_name}")
            run_s2.font.name = CFG.font_name
            run_s2.font.size = Pt(9)
            run_s2.bold = True

            doc.add_paragraph("")

    # Signature
    add_section_title(doc, "Потпис")
    doc.add_paragraph(f"Дефектолог: ____________________          Родител/старател: ____________________")

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"RODITELSKI_SREDBI_{school_year}.docx"
    doc.save(out_path)
    return out_path


# ----------------------------
# Родителски состаноци generator
# ----------------------------
def generate_roditelski_sostanoci(
    out_dir: Path,
    students: List[Dict[str, Any]],
    student_names: Dict[str, str],
    diary: Dict[str, Any],
    school_year: str = "2025-2026",
) -> Path:
    """Generate parent group meeting documents (sostanoci)."""
    doc = Document()
    apply_default_styles(doc, font_name=CFG.font_name, font_size=CFG.font_size)

    mk_header(doc)
    add_section_title(doc, "РОДИТЕЛСКИ СОСТАНОЦИ", level=1)
    p = doc.add_paragraph(school_year)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"Одделение: ___________       Терапевт: {CFG.therapist_name}")
    doc.add_paragraph("")

    # 5 meetings based on annual plan
    meetings = [
        (f"{school_year[:4]}-09-18", [
            "Разговори за состојбата и тековни потреби на учениците",
            "Запознавање со планот и програмата за рехабилитација на говор и јазик",
            "Резултати од претходни испитувања (аудиограми, тријажни тестови) и почетни насоки",
            "Консултации за чување и редовно носење на слушни апарати/помагала",
            "Насоки за домашни услови: комуникација, рутини и поддршка",
            "Разно / прашања",
        ], "Воведно информирање и договор за соработка со родители за учебната година."),
        (f"{school_year[:4]}-11-27", [
            "Известување за реализираните третмани во изминатиот период",
            "Резултати, забелешки и примери на напредок",
            "Разговор за тековни проблеми и предлози за надминување",
            "Консултации за чување и редовно носење на апарати/помагала",
            "План за активности и домашни насоки за наредниот период",
            "Разно / прашања",
        ], "Информација за реализирани третмани и првични резултати; договор за следниот период."),
        (f"{school_year[-4:]}-01-29", [
            "Полугодишен преглед на напредокот и постигнатите резултати",
            "Начинот на прифаќање на планираните вежби и терапевтски активности",
            "Темпо на работа и напредување: индивидуални разлики",
            "Насоки за родителите на што да се обрнува внимание во домашни услови",
            "Консултации за апарати/помагала (проверка, одржување, сервис)",
            "Разно / прашања",
        ], "Полугодишен преглед: темпо на работа, прифаќање на вежби и насоки за домашни услови."),
        (f"{school_year[-4:]}-04-23", [
            "Разговори за состојбата и тековни проблеми со учениците",
            "Континуитет и напредок во терапевтските активности во третото тромесечие",
            "Како детето ги прифаќа задачите: мотивација, внимание, соработка",
            "Насоки за домашни услови и поддршка (говор, слух, комуникација)",
            "Планирање на активности за летниот период и премин во следната година",
            "Разно / прашања",
        ], "Пролетен состанок: континуитет, напредок и подготовка за завршување на годината."),
        (f"{school_year[-4:]}-06-11", [
            "Известување за реализацијата на планот и програмата за учебната година",
            "Континуитетот и напредокот во работата: реализација и постигања",
            "Планирања за наредната учебна година (рамка на цели и динамика)",
            "Упатства и насоки за родителите за време на летниот распуст",
            "Консултации за апарати/помагала и препораки за контроли",
            "Разно / прашања",
        ], "Завршен годишен состанок: реализација, напредок и насоки за летниот распуст."),
    ]

    # Collect all parent names for signatures
    parent_names = []
    targets = sorted(student_names.keys(), key=lambda x: normalize_whitespace(x))
    for nm in targets:
        sid = student_names.get(nm)
        s = next((x for x in students if str(x.get("id")) == sid), None)
        if not s:
            continue
        father = (s.get("fatherName") or "").strip()
        mother = (s.get("motherName") or "").strip()
        if father and father not in parent_names:
            parent_names.append(father)
        if mother and mother not in parent_names:
            parent_names.append(mother)

    for i, (date, agenda, notes) in enumerate(meetings, 1):
        # Meeting box
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]

        p_title = cell.paragraphs[0]
        run = p_title.add_run(f"{i}. РОДИТЕЛСКИ СОСТАНОК    Дата: {date}")
        run.bold = True
        run.font.name = CFG.font_name
        run.font.size = Pt(CFG.font_size)

        p_agenda = cell.add_paragraph()
        run_a = p_agenda.add_run("Дневен ред:")
        run_a.bold = True
        run_a.font.name = CFG.font_name

        for j, item in enumerate(agenda, 1):
            p_item = cell.add_paragraph(f"    {j}. {item}")
            for r in p_item.runs:
                r.font.name = CFG.font_name
                r.font.size = Pt(CFG.font_size)

        cell.add_paragraph("")

        p_sig_title = cell.add_paragraph()
        run_st = p_sig_title.add_run("Потписи на родители:")
        run_st.bold = True
        run_st.font.name = CFG.font_name

        for j, pname in enumerate(parent_names[:10], 1):
            p_pn = cell.add_paragraph(f"    {j}. {pname}: ________________")
            for r in p_pn.runs:
                r.font.name = CFG.font_name
                r.font.size = Pt(CFG.font_size)

        cell.add_paragraph("")
        p_notes = cell.add_paragraph()
        run_n = p_notes.add_run("Забелешки:")
        run_n.bold = True
        run_n.font.name = CFG.font_name
        cell.add_paragraph(notes)

        cell.add_paragraph("")
        p_therapist = cell.add_paragraph()
        p_therapist.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_th = p_therapist.add_run(f"Терапевт:\n{CFG.therapist_name}")
        run_th.font.name = CFG.font_name

        doc.add_paragraph("")

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"RODITELSKI_SOSTANOCI_{school_year}.docx"
    doc.save(out_path)
    return out_path


# ----------------------------
# Извештај (сумирано) generator
# ----------------------------
def generate_izvestaj(
    out_dir: Path,
    students: List[Dict[str, Any]],
    student_names: Dict[str, str],
    diary: Dict[str, Any],
    triage_by_sid,
    audi_by_name,
    aliases: Dict[str, str],
    school_year: str = "2025-2026",
) -> Path:
    """Generate summary report (Извештај сумирано) with all students."""
    doc = Document()
    apply_default_styles(doc, font_name=CFG.font_name, font_size=CFG.font_size)
    mk_header(doc)

    add_section_title(doc, f"ИЗВЕШТАЈ (СУМИРАНО) – {school_year}", level=1)
    doc.add_paragraph("Извештајот содржи сумирање на клучни информации од досиејата и евиденциите по ученици.")
    doc.add_paragraph("")

    # Summary table
    add_section_title(doc, "1. Преглед по ученици", level=2)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Ученик", "Оддел.", "План", "Сесии", "Присуство %", "Тријажа", "Аудиогр."]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = CFG.font_name

    low_attendance = []
    targets = sorted(student_names.keys(), key=lambda x: normalize_whitespace(x))

    for nm in targets:
        sid = student_names.get(nm)
        s = next((x for x in students if str(x.get("id")) == sid), None)
        if not s:
            continue

        # Plan
        diary_student = next((x for x in (diary.get("students") or []) if str(x.get("id")) == sid), None)
        plan_name = "—"
        if diary_student:
            pid = diary_student.get("planId")
            plan = next((p for p in (diary.get("plans") or []) if p.get("id") == pid), None)
            if plan:
                plan_name = plan.get("name", "—")

        # Sessions count
        sp = (diary.get("studentProgress") or {}).get(sid, {}) or {}
        sessions = sum(len(v) for v in sp.values() if isinstance(v, list))

        # Attendance
        present = absent = 0
        for _, by_student in (diary.get("attendance") or {}).items():
            st = by_student.get(sid) if isinstance(by_student, dict) else None
            if isinstance(st, dict):
                for _, status in st.items():
                    if status == "present":
                        present += 1
                    elif status == "absent":
                        absent += 1
        total = present + absent
        pct = (present / total * 100) if total else None
        pct_str = f"{pct:.1f}%" if pct is not None else "—"

        if pct is not None and pct < 70:
            low_attendance.append((nm, pct))

        # Triage
        tri = triage_by_sid.get(sid, [])
        tri_str = safe_date_fmt(tri[0].get("date", "")) if tri else "—"

        # Audiograms
        name_norm = normalize_whitespace(nm)
        name_for_audi = name_norm if name_norm in audi_by_name else normalize_whitespace(aliases.get(nm, nm))
        audi_count = len(audi_by_name.get(name_for_audi, []))

        row = table.add_row().cells
        values = [nm, s.get("grade", ""), plan_name, str(sessions), pct_str, tri_str, str(audi_count) if audi_count else "—"]
        for i, v in enumerate(values):
            row[i].text = v
            for run in row[i].paragraphs[0].runs:
                run.font.name = CFG.font_name
                run.font.size = Pt(10)

    doc.add_paragraph("")

    # Low attendance notes
    add_section_title(doc, "2. Забелешки за отстапувања / специфични ситуации", level=2)
    if low_attendance:
        for nm, pct in low_attendance:
            doc.add_paragraph(f"- {nm}: ниско присуство ({pct:.1f}%).")
    else:
        doc.add_paragraph("Нема забележани отстапувања.")
    doc.add_paragraph("")

    add_section_title(doc, "3. Заклучок", level=2)
    doc.add_paragraph("Овој документ е автоматски генериран како нацрт и служи за побрза подготовка. Финалниот текст се ревидира и дополнува рачно.")
    doc.add_paragraph("")

    add_section_title(doc, "Потпис")
    doc.add_paragraph(f"{CFG.therapist_title} {CFG.therapist_name}: ____________________    Датум: ______________")

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"IZVESTAJ_Sumirano_{school_year}_AUTO.docx"
    doc.save(out_path)
    return out_path


# ----------------------------
# Периодичен извештај generator
# ----------------------------
def generate_periodicen_izvestaj(
    out_dir: Path,
    students: List[Dict[str, Any]],
    student_names: Dict[str, str],
    diary: Dict[str, Any],
    triage_by_sid,
    audi_by_name,
    aliases: Dict[str, str],
    school_year: str = "2025-2026",
    period_label: str = "",
) -> Path:
    """Generate periodic report (Периодичен извештај) - narrative per student grouped by grade.

    Matches the template structure:
    - Header: Cabinet name centered, report period centered
    - Body: Students grouped by одделение, each with narrative paragraph
    - Style: Times New Roman 12pt, bold for student paragraphs
    """
    doc = Document()
    apply_default_styles(doc, font_name=CFG.font_name, font_size=CFG.font_size)

    # Header - centered institution info
    mk_header(doc)
    p_cab = doc.add_paragraph(CFG.cabinet)
    p_cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_cab.runs:
        r.bold = True
        r.font.name = CFG.font_name
        r.font.size = Pt(14)
    doc.add_paragraph("")

    # Period label
    if not period_label:
        period_label = f"Извештај за учебна {school_year}"
    p_period = doc.add_paragraph(period_label)
    p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_period.runs:
        r.bold = True
        r.font.name = CFG.font_name
        r.font.size = Pt(CFG.font_size)
    doc.add_paragraph("")

    # Group students by grade
    grade_groups = defaultdict(list)
    targets = sorted(student_names.keys(), key=lambda x: normalize_whitespace(x))

    for nm in targets:
        sid = student_names.get(nm)
        s = next((x for x in students if str(x.get("id")) == sid), None)
        if not s:
            continue
        grade = (s.get("grade") or "").strip() or "(без одделение)"
        grade_groups[grade].append((nm, s, sid))

    # Sort grades (roman numeral order approximation)
    def grade_sort_key(g):
        g_upper = g.upper().replace("-", "").replace(" ", "")
        roman_map = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8, "IX": 9}
        for roman, num in sorted(roman_map.items(), key=lambda x: -x[1]):
            if g_upper.startswith(roman):
                return (num, g)
        if g.startswith("("):
            return (99, g)
        return (50, g)

    sorted_grades = sorted(grade_groups.keys(), key=grade_sort_key)

    for grade in sorted_grades:
        students_in_grade = grade_groups[grade]

        # Grade heading (bold, like template)
        add_section_title(doc, f"{grade} одделение", level=2)

        for nm, s, sid in students_in_grade:
            # Build narrative paragraph for this student
            # Pull diagnosis from findings
            findings = (s.get("findings") or "").strip()
            opinion = (s.get("opinion") or "").strip()

            # Get session count
            sp = (diary.get("studentProgress") or {}).get(sid, {}) or {}
            sessions = sum(len(v) for v in sp.values() if isinstance(v, list))

            # Attendance
            present = absent = 0
            for _, by_student in (diary.get("attendance") or {}).items():
                st = by_student.get(sid) if isinstance(by_student, dict) else None
                if isinstance(st, dict):
                    for _, status in st.items():
                        if status == "present":
                            present += 1
                        elif status == "absent":
                            absent += 1
            total = present + absent
            pct = (present / total * 100) if total else None

            # Build narrative (similar to how the template reads)
            narrative_parts = []

            # Student name + diagnosis summary
            if findings:
                # Extract key diagnosis info
                diag_short = findings.split("\n")[0].strip()
                if len(diag_short) > 120:
                    diag_short = diag_short[:120] + "..."
                narrative_parts.append(f"{nm} - {diag_short}")
            else:
                narrative_parts.append(f"{nm}")

            # Add progress note
            if sessions > 0:
                narrative_parts.append(f"Реализирани {sessions} сесии.")
            if pct is not None:
                if pct >= 80:
                    narrative_parts.append(f"Редовно присуство ({pct:.0f}%).")
                elif pct >= 50:
                    narrative_parts.append(f"Присуство {pct:.0f}%.")
                else:
                    narrative_parts.append(f"Ниско присуство ({pct:.0f}%).")

            narrative_parts.append("Соработува на третманите и постигнува напредок според способностите и можностите.")

            # Write as single bold paragraph (matching template style)
            text = " ".join(narrative_parts)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = CFG.font_name
            run.font.size = Pt(CFG.font_size)

        doc.add_paragraph("")

    # Signature
    doc.add_paragraph("")
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_s = p_sign.add_run(f"{CFG.therapist_title}\n{CFG.therapist_name}")
    run_s.bold = True
    run_s.font.name = CFG.font_name
    run_s.font.size = Pt(CFG.font_size)

    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = period_label.replace(" ", "_").replace("/", "-").replace(".", "")[:50] if period_label else "Periodicen"
    out_path = out_dir / f"IZVESTAJ_{safe_name}_{school_year}.docx"
    doc.save(out_path)
    return out_path


# ----------------------------
# CLI
# ----------------------------
def find_latest(pattern: str, base: Path) -> Optional[Path]:
    files = sorted(base.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def find_latest_any(base: Path, patterns: List[str]) -> Optional[Path]:
    files: List[Path] = []
    for pattern in patterns:
        files.extend(base.glob(pattern))
    files = [p for p in files if p.exists() and p.is_file()]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def safe_print(msg: str):
    """Print safely on Windows consoles (no UnicodeEncodeError)."""
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("utf-8", errors="replace").decode("utf-8", errors="replace"))


def main():
    # Fix Windows console encoding
    if hasattr(sys.stdout, "reconfigure"):
        try:
            sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass

    ap = argparse.ArgumentParser(description="SMUROP Document Factory v4.0")
    ap.add_argument("--student-db", type=str, default=None, help="Path to student_db_backup_*.json")
    ap.add_argument("--diary", type=str, default=None, help="Path to diary_export_*.json")
    ap.add_argument("--audi", type=str, default=None, help="Path to AUDIOGRAMI_BAZA_*.json")
    ap.add_argument("--unified", type=str, default=None, help="Path to unified e_dnevnik state JSON")
    ap.add_argument("--overrides", type=str, default=None, help="Path to PROFILE_OVERRIDES.json (optional)")
    ap.add_argument("--student-patch", type=str, default=None, help="Path to student_records_patch.json (optional)")
    ap.add_argument("--audi-patch", type=str, default=None, help="Path to audiogram_images_patch.json (optional)")
    ap.add_argument("--out", type=str, default="OUT", help="Output folder")
    ap.add_argument("--school-year", type=str, default="2025-2026")
    ap.add_argument("--font-name", type=str, default="Times New Roman", help="DOCX font name (Cyrillic-safe)")
    ap.add_argument("--font-size", type=int, default=12, help="Base font size (default 12)")
    ap.add_argument("--student", type=str, default=None, help="Generate only for this full name (exact)")
    ap.add_argument("--images-folder", type=str, default=None, help="Folder where image files live (optional)")
    ap.add_argument("--download-images", action="store_true", help="Try to download direct image URLs and embed them")
    # v4 new args
    ap.add_argument("--doc-types", type=str, default="dosie",
                     help="Comma-separated: dosie,roditelski,sostanoci,izvestaj,periodicen (default: dosie)")
    ap.add_argument("--output-mode", type=str, default="separate", choices=["separate", "merged", "both"],
                     help="separate=one file/student, merged=single file, both=both")
    ap.add_argument("--therapist-name", type=str, default=None, help="Therapist name (default: Благој Насев)")
    ap.add_argument("--therapist-title", type=str, default=None, help="Therapist title")
    ap.add_argument("--institution", type=str, default=None, help="Institution full name")
    ap.add_argument("--period-label", type=str, default="", help="Label for periodic report (e.g. 'Извештај 01-07.10.2025')")

    args = ap.parse_args()

    # Apply config
    CFG.font_name = args.font_name
    CFG.font_size = args.font_size
    CFG.school_year = args.school_year
    if args.therapist_name:
        CFG.therapist_name = args.therapist_name
    if args.therapist_title:
        CFG.therapist_title = args.therapist_title
    if args.institution:
        CFG.institution = args.institution

    base = Path(__file__).resolve().parent

    student_db_path = Path(args.student_db) if args.student_db else find_latest("student_db_backup_*.json", base)
    diary_path = Path(args.diary) if args.diary else find_latest("diary_export_*.json", base)
    audi_path = Path(args.audi) if args.audi else find_latest("AUDIOGRAMI_BAZA_*.json", base)
    unified_path = Path(args.unified) if args.unified else find_latest_any(base, [
        "e_dnevnik_unified_state*.json",
        "*unified_state*.json",
        "*unified*.json",
    ])
    overrides_path = Path(args.overrides) if args.overrides else (base / "PROFILE_OVERRIDES.json")

    use_unified = False
    if args.unified:
        if not unified_path or not unified_path.exists():
            raise SystemExit("Missing unified JSON path passed to --unified.")
        use_unified = True
    else:
        split_ready = bool(student_db_path and student_db_path.exists() and diary_path and diary_path.exists())
        if split_ready:
            use_unified = False
        elif unified_path and unified_path.exists():
            use_unified = True
        else:
            raise SystemExit(
                "Missing inputs. Provide split files (student_db_backup + diary_export) or a unified JSON (--unified)."
            )

    if use_unified:
        safe_print(f"[INFO] Loading unified state: {unified_path}")
        unified_payload = load_json(unified_path)
        student_db = {
            "meta": {"source": "unified"},
            "data": {"student_records": extract_student_records(unified_payload)},
        }
        diary = extract_diary_payload(unified_payload)
        audi_db = extract_audi_records(unified_payload)
        audi_path = unified_path
    else:
        safe_print(f"[INFO] Loading student DB: {student_db_path}")
        student_db = load_json(student_db_path)
        safe_print(f"[INFO] Loading diary: {diary_path}")
        diary = load_json(diary_path)
        if not audi_path or not audi_path.exists():
            audi_db = []
        else:
            audi_db = normalize_audi_db(load_json(audi_path))

    overrides = {}
    if overrides_path.exists():
        overrides = load_json(overrides_path)

    # Apply patches (optional)
    if args.student_patch:
        sp = load_json(Path(args.student_patch))
        applied = apply_student_patch(student_db, sp)
        safe_print(f"[INFO] Applied student patch: {applied} records")
    if args.audi_patch and audi_path and audi_path.exists():
        apatch_raw = load_json(Path(args.audi_patch))
        apatch = parse_audi_patch(apatch_raw)
        applied = apply_audi_image_patch(audi_db, apatch)
        safe_print(f"[INFO] Applied audiogram image patch: {applied} records")

    out_dir = Path(args.out)
    images_folder = Path(args.images_folder) if args.images_folder else None

    students, student_names, triage_by_sid, audi_by_name, aliases, categories = build_indexes(student_db, diary, audi_db, overrides)

    safe_print(f"[INFO] Found {len(students)} student(s)")

    doc_types = [t.strip().lower() for t in args.doc_types.split(",") if t.strip()]
    safe_print(f"[INFO] Document types: {doc_types}")
    safe_print(f"[INFO] Output mode: {args.output_mode}")

    results = []

    # --- DOSIE (per student) ---
    if "dosie" in doc_types:
        targets = []
        if args.student:
            targets = [args.student]
        else:
            targets = sorted(student_names.keys(), key=lambda x: normalize_whitespace(x))

        safe_print(f"\n-- Generating DOSIE for {len(targets)} student(s) --")
        for nm in targets:
            try:
                p = generate_dossier(
                    out_dir=out_dir,
                    students=students,
                    student_names=student_names,
                    diary=diary,
                    triage_by_sid=triage_by_sid,
                    audi_by_name=audi_by_name,
                    aliases=aliases,
                    categories=categories,
                    full_name=nm,
                    school_year=args.school_year,
                    images_folder=images_folder,
                    download_images=args.download_images,
                )
                safe_print(f"  OK: {p.name}")
                results.append(p)
            except Exception as e:
                safe_print(f"  SKIP: {nm} - {e}")

    # --- RODITELSKI SREDBI (individual parent meetings) ---
    if "roditelski" in doc_types:
        safe_print("\n-- Generating RODITELSKI SREDBI --")
        try:
            p = generate_roditelski(
                out_dir=out_dir,
                students=students,
                student_names=student_names,
                diary=diary,
                school_year=args.school_year,
            )
            safe_print(f"  OK: {p.name}")
            results.append(p)
        except Exception as e:
            safe_print(f"  FAIL: {e}")

    # --- RODITELSKI SOSTANOCI (group parent meetings) ---
    if "sostanoci" in doc_types:
        safe_print("\n-- Generating RODITELSKI SOSTANOCI --")
        try:
            p = generate_roditelski_sostanoci(
                out_dir=out_dir,
                students=students,
                student_names=student_names,
                diary=diary,
                school_year=args.school_year,
            )
            safe_print(f"  OK: {p.name}")
            results.append(p)
        except Exception as e:
            safe_print(f"  FAIL: {e}")

    # --- IZVESTAJ (summary report) ---
    if "izvestaj" in doc_types:
        safe_print("\n-- Generating IZVESTAJ (sumirano) --")
        try:
            p = generate_izvestaj(
                out_dir=out_dir,
                students=students,
                student_names=student_names,
                diary=diary,
                triage_by_sid=triage_by_sid,
                audi_by_name=audi_by_name,
                aliases=aliases,
                school_year=args.school_year,
            )
            safe_print(f"  OK: {p.name}")
            results.append(p)
        except Exception as e:
            safe_print(f"  FAIL: {e}")

    # --- PERIODICEN IZVESTAJ (narrative periodic report per grade) ---
    if "periodicen" in doc_types:
        safe_print("\n-- Generating PERIODICEN IZVESTAJ --")
        try:
            p = generate_periodicen_izvestaj(
                out_dir=out_dir,
                students=students,
                student_names=student_names,
                diary=diary,
                triage_by_sid=triage_by_sid,
                audi_by_name=audi_by_name,
                aliases=aliases,
                school_year=args.school_year,
                period_label=args.period_label,
            )
            safe_print(f"  OK: {p.name}")
            results.append(p)
        except Exception as e:
            safe_print(f"  FAIL: {e}")

    safe_print(f"\n[DONE] Generated {len(results)} file(s) in: {out_dir.resolve()}")


if __name__ == "__main__":
    main()
