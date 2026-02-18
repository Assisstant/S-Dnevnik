#!/usr/bin/env python3
"""
Compact DOCX generator from unified JSON + DOCX template.

Usage:
  python compact_factory.py --unified e_dnevnik_unified_state_v7.json --template template.docx --out Dosijea

Template placeholders: use double-curly keys like {{firstName}} {{lastName}} {{grade}} etc.
This tool performs simple textual replacement. If a field contains image links those are
written as text markers ("[image: ...]") rather than embedded to keep the script compact
and robust across environments.
"""
import argparse
import datetime
import io
import json
import re
import shutil
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt
import sys
import io

# Ensure stdout/stderr use UTF-8 so printing Unicode filenames won't fail on Windows
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass
else:
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    except Exception:
        pass


def load_json(p: Path) -> Any:
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)


def safe_date_fmt(iso: str) -> str:
    if not iso:
        return ""
    try:
        return datetime.date.fromisoformat(iso).strftime("%d.%m.%Y")
    except Exception:
        return iso


def extract_students_from_unified(payload: Any) -> List[Dict[str, Any]]:
    if isinstance(payload, dict):
        for key in ("students", "student_records", "studentRecords"):
            arr = payload.get(key)
            if isinstance(arr, list):
                return [normalize_student_record(x, i + 1) for i, x in enumerate(arr)]
        if isinstance(payload.get("data"), dict) and isinstance(payload["data"].get("student_records"), list):
            return [normalize_student_record(x, i + 1) for i, x in enumerate(payload["data"]["student_records"])]
    if isinstance(payload, list):
        return [normalize_student_record(x, i + 1) for i, x in enumerate(payload)]
    return []


def extract_json_from_html(html_path: Path, out_json: Path) -> Optional[Path]:
    """Try to find a JSON object containing 'students' inside an HTML file and save it."""
    s = html_path.read_text(encoding="utf-8", errors="ignore")
    # find a position of the word "students"
    m = re.search(r"\"students\"", s)
    if not m:
        m = re.search(r"students\s*:", s)
    if not m:
        return None
    pos = m.start()

    # find opening brace before pos
    start = s.rfind("{", 0, pos)
    if start == -1:
        start = s.find("{", 0, pos)
        if start == -1:
            return None

    # scan forward to match braces
    depth = 0
    end = -1
    for i in range(start, len(s)):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break

    if end == -1:
        return None

    candidate = s[start:end]
    # try to load JSON
    try:
        obj = json.loads(candidate)
        out_json.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
        return out_json
    except Exception:
        # fallback: try to locate any JSON-like {...} that contains "students"
        for jm in re.finditer(r"\{.*?\}", s, flags=re.DOTALL):
            chunk = jm.group(0)
            if "\"students\"" in chunk or "students" in chunk:
                try:
                    obj = json.loads(chunk)
                    out_json.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
                    return out_json
                except Exception:
                    continue
    return None


def list_students_in_unified(unified_path: Path) -> List[str]:
    payload = json.loads(unified_path.read_text(encoding="utf-8"))
    students = extract_students_from_unified(payload)
    names = [f"{s.get('firstName','').strip()} {s.get('lastName','').strip()}".strip() for s in students]
    return names


def assemble_final_folder(source_dir: Path, template_paths: List[Path], generated_dir: Path, final_dir: Path, html_path: Optional[Path] = None):
    final_dir.mkdir(parents=True, exist_ok=True)
    # copy S-Dnevnik.html if provided
    if html_path and html_path.exists():
        shutil.copy2(html_path, final_dir / html_path.name)
    # copy json files from source_dir
    for p in source_dir.glob("*.json"):
        shutil.copy2(p, final_dir / p.name)
    # copy templates
    for t in template_paths:
        if t.exists():
            shutil.copy2(t, final_dir / t.name)
    # copy generated docs
    if generated_dir.exists():
        for g in generated_dir.glob("*.docx"):
            shutil.copy2(g, final_dir / g.name)
    return final_dir


def normalize_student_record(raw: Dict[str, Any], fallback_id: int = 0) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return {}
    rec = dict(raw)
    rec.setdefault("firstName", (rec.get("firstName") or ""))
    rec.setdefault("lastName", (rec.get("lastName") or ""))
    if not rec.get("firstName") and rec.get("name"):
        parts = str(rec.get("name") or "").split(" ", 1)
        rec["firstName"] = parts[0] if parts else ""
        rec["lastName"] = parts[1] if len(parts) > 1 else rec.get("lastName", "")
    rid = rec.get("id") or rec.get("studentId") or fallback_id
    rec["id"] = rid
    rec.setdefault("grade", "")
    rec.setdefault("birthDate", "")
    rec.setdefault("contact", "")
    rec.setdefault("fatherName", "")
    rec.setdefault("motherName", "")
    rec.setdefault("address", "")
    rec.setdefault("residence", "")
    rec.setdefault("findings", "")
    rec.setdefault("opinion", "")
    rec.setdefault("attachmentLinks", [])
    return rec


def flatten_student(rec: Dict[str, Any]) -> Dict[str, str]:
    out: Dict[str, str] = {}
    out["id"] = str(rec.get("id", ""))
    out["firstName"] = rec.get("firstName", "")
    out["lastName"] = rec.get("lastName", "")
    out["name"] = f"{rec.get('firstName','')} {rec.get('lastName','')}".strip()
    out["grade"] = rec.get("grade", "")
    out["birthDate"] = safe_date_fmt(rec.get("birthDate", ""))
    out["contact"] = rec.get("contact", "")
    out["fatherName"] = rec.get("fatherName", "")
    out["motherName"] = rec.get("motherName", "")
    out["address"] = rec.get("address", "")
    out["residence"] = rec.get("residence", "")
    out["findings"] = rec.get("findings", "")
    out["opinion"] = rec.get("opinion", "")
    links = rec.get("attachmentLinks") or []
    if isinstance(links, list):
        out["attachments"] = "; ".join([str(x) for x in links])
    else:
        out["attachments"] = str(links or "")
    return out


def replace_in_paragraph(par, mapping: Dict[str, str]):
    text = par.text or ""
    if not text:
        return
    replaced = text
    for k, v in mapping.items():
        token = "{{" + k + "}}"
        if token in replaced:
            # If value looks like an image link, write a marker instead of embedding
            if isinstance(v, str) and (v.startswith("data:image/") or v.lower().endswith(('.png', '.jpg', '.jpeg', '.webp'))):
                replaced = replaced.replace(token, f"[image: {v}]")
            else:
                replaced = replaced.replace(token, str(v))
    if replaced != text:
        # collapse runs and set single run with replaced text
        for r in list(par.runs):
            r.clear()
        par.add_run(replaced)


def replace_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                replace_in_paragraph(par, mapping)


def replace_in_doc(doc: Document, mapping: Dict[str, str]):
    for par in doc.paragraphs:
        replace_in_paragraph(par, mapping)
    for table in doc.tables:
        replace_in_table(table, mapping)
    for section in doc.sections:
        for par in section.header.paragraphs:
            replace_in_paragraph(par, mapping)
        for par in section.footer.paragraphs:
            replace_in_paragraph(par, mapping)


def cmd_extract(args):
    html = Path(args.html)
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    target = out / (args.json_name or (html.stem + ".json"))
    res = extract_json_from_html(html, target)
    if res:
        print(f"Extracted JSON to: {res}")
    else:
        print("Failed to extract JSON from HTML.")


def cmd_list(args):
    unified = Path(args.unified)
    names = list_students_in_unified(unified)
    for n in names:
        print(n)


def cmd_gen_template(args):
    unified = load_json(Path(args.unified))
    students = extract_students_from_unified(unified)
    if not students:
        print("No students found in unified JSON")
        return
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    for s in students:
        full = f"{s.get('firstName','').strip()} {s.get('lastName','').strip()}".strip()
        if args.student and args.student != full:
            continue
        mapping = flatten_student(s)
        doc = Document(Path(args.template))
        replace_in_doc(doc, mapping)
        safe_name = full.replace(' ', '_') or mapping.get('id', 'unknown')
        out_path = out_dir / f"DOC_{safe_name}.docx"
        doc.save(out_path)
        print(f"Wrote: {out_path}")


def cmd_assemble(args):
    src = Path(args.source_dir)
    final = Path(args.final_dir)
    template_list = [Path(p) for p in (args.templates or [])]
    gen_dir = Path(args.generated_dir or "OUT")
    html_path = Path(args.html) if args.html else None
    res = assemble_final_folder(src, template_list, gen_dir, final, html_path)
    print(f"Assembled final folder: {res}")


def main():
    ap = argparse.ArgumentParser(description="Compact DOCX generator / workflow helper")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p_extract = sub.add_parser("extract", help="Extract unified JSON from S-Dnevnik HTML")
    p_extract.add_argument("--html", required=True, help="Path to S-Dnevnik.html")
    p_extract.add_argument("--out", default=".", help="Folder to write JSON")
    p_extract.add_argument("--json-name", default=None, help="Output json filename (optional)")
    p_extract.set_defaults(func=cmd_extract)

    p_list = sub.add_parser("list-students", help="List students in a unified JSON")
    p_list.add_argument("--unified", required=True, help="Path to unified JSON")
    p_list.set_defaults(func=cmd_list)

    p_gen = sub.add_parser("gen-template", help="Generate DOCX files from template and unified JSON")
    p_gen.add_argument("--unified", required=True)
    p_gen.add_argument("--template", required=True, help="DOCX template with {{placeholders}}")
    p_gen.add_argument("--out", default="OUT")
    p_gen.add_argument("--student", default=None, help="Full name to generate only for one student (exact)")
    p_gen.set_defaults(func=cmd_gen_template)

    p_assemble = sub.add_parser("assemble-final", help="Create a final/ folder with html, json and generated docs")
    p_assemble.add_argument("--source-dir", default=".", help="Folder with exported JSONs")
    p_assemble.add_argument("--templates", nargs="*", help="List of template files to include")
    p_assemble.add_argument("--generated-dir", default="OUT", help="Folder where generated docs live")
    p_assemble.add_argument("--final-dir", default="final", help="Destination final folder")
    p_assemble.add_argument("--html", default=None, help="Optional path to S-Dnevnik.html to include")
    p_assemble.set_defaults(func=cmd_assemble)

    args = ap.parse_args()
    args.func(args)


if __name__ == '__main__':
    main()
